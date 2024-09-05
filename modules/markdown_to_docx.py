import markdown
from docx import Document
from bs4 import BeautifulSoup

def convert_markdown_to_docx(markdown_string, output_file):
    """Convert a markdown string to a docx file.

    Args:
        markdown_string (str): The markdown string to convert.
        output_file (str): The output docx file.
    """
    
    # Convert markdown to HTML
    html_string = markdown.markdown(markdown_string)
    
    # Parse the HTML content
    soup = BeautifulSoup(html_string, 'html.parser')
    
    # Create a new Document
    doc = Document()
    
    # Convert HTML to docx
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name in ['strong', 'b']:
            p = doc.add_paragraph()
            p.add_run(element.text).bold = True
        elif element.name in ['em', 'i']:
            p = doc.add_paragraph()
            p.add_run(element.text).italic = True
    
    # Save the document
    doc.save(output_file)
