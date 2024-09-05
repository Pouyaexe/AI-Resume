

import markdown
from docx import Document
from bs4 import BeautifulSoup

def convert_markdown_to_docx(markdown_string, output_file):
    """Convert a markdown string to a docx file.

    Args:
        markdown_string (str): The markdown string to convert.
        output_file (str): The output docx file.
    """
    
    # Convert markdown to HTML
    html_string = markdown.markdown(markdown_string)
    
    # Parse the HTML content
    soup = BeautifulSoup(html_string, 'html.parser')
    
    # Create a new Document
    doc = Document()
    
    # Convert HTML to docx
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name in ['strong', 'b']:
            p = doc.add_paragraph()
            p.add_run(element.text).bold = True
        elif element.name in ['em', 'i']:
            p = doc.add_paragraph()
            p.add_run(element.text).italic = True
    
    # Save the document
    doc.save(output_file)

markdown_text = """
## John Doe

123 Main Street, City, Country | john.doe@email.com | (123) 456-7890

**Education**

* **Bachelor of Arts, History, XYZ University, May 2020**
  * Dean's List, 2018-2020 (Optional, if applicable)
  * Relevant Coursework: Research Methods, Historical Analysis, Public Speaking (Optional, if applicable)
  * Completed a research project on the history of... (Optional, if applicable)

**Work Experience**

**Customer Service Representative**
ABC Retail, June 2021 – Present

* Assisted over 100 customers daily with purchases, resolving issues and ensuring satisfaction.
* Managed checkout operations efficiently, processing transactions and providing accurate change.
* Assisted in maintaining accurate product inventory by performing regular stock checks and replenishing displays.

**Intern**
XYZ Museum, Summer 2019

* Collaborated with the exhibits team to install and maintain engaging displays, ensuring a positive visitor experience.
* Conducted informative tours for museum visitors, sharing knowledge of exhibits and museum history.
* Supported the marketing team in promoting museum events, contributing to increased visitor attendance.

**Skills**

* **Communication:**
  * Presented complex technical information to non-technical audiences, resulting in a 15% increase in project understanding. (Optional, if applicable)
  * Collaborated with diverse teams across multiple departments to facilitate successful project launches. (Optional, if applicable)
* **Time Management:**
  * Prioritized multiple tasks and deadlines effectively, consistently delivering projects on time and within budget. (Optional, if applicable)
  * Managed a team of 5 individuals, streamlining workflow and improving project delivery efficiency by 20%. (Optional, if applicable)
* **Microsoft Office:**
  * Proficient in Microsoft Word, Excel, PowerPoint, and Outlook.
  * Developed and maintained complex spreadsheets in Excel, automating data analysis and reporting. (Optional, if applicable)

**Additional Skills:**

* Developed critical analysis skills by analyzing historical narratives and identifying key themes and events.
* Enhanced research abilities by researching historical periods and figures.
* Demonstrated teamwork and leadership skills as a member of a competitive soccer team.
* Developed strong communication and coordination skills while collaborating with teammates.
"""

output_file = "output111.docx"
convert_markdown_to_docx(markdown_text,
                         output_file)
